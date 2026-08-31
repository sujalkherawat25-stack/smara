import os
import json
import subprocess
from pathlib import Path

import pytest

from smara.desktop_executor import (
    DesktopRunner,
    ExecutionCancelled,
    _load_state,
    _save_state,
    delete_local_credential,
    execute_step,
    local_connector_audit,
    local_connector_summaries,
    local_credential_summaries,
    revoke_local_connector,
    resolve_local_credential,
    save_local_credential,
)
from smara.desktop_integrations import local_connector_catalog


def test_desktop_file_read_write_stays_inside_approved_root(tmp_path: Path):
    state = {"capabilities": ["local_file_read", "local_file_write"], "allowed_roots": [str(tmp_path)]}
    source = tmp_path / "note.txt"
    source.write_text("hello", encoding="utf-8")
    read = json.loads(execute_step({"required_capability": "local_file_read", "executor_payload": {"path": str(source)}}, state))
    assert read["sha256"]
    assert read["content_shared"] is False
    written = json.loads(execute_step({"required_capability": "local_file_write", "executor_payload": {"path": str(tmp_path / "out.txt"), "content": "done"}}, state))
    assert written["bytes_written"] == 4
    assert (tmp_path / "out.txt").read_text(encoding="utf-8") == "done"
    with pytest.raises(RuntimeError, match="outside"):
        execute_step({"required_capability": "local_file_read", "executor_payload": {"path": str(tmp_path.parent / "note.txt")}}, state)


def test_desktop_file_preview_is_read_only_and_write_returns_diff_and_undo(tmp_path: Path):
    target = tmp_path / "note.txt"
    target.write_text("before\n", encoding="utf-8")
    state = {
        "capabilities": ["local_file_write"],
        "allowed_roots": [str(tmp_path)],
        "_state_path": str(tmp_path / "desktop.json"),
    }
    preview = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "write", "path": str(target), "content": "after\n", "preview_only": True},
    }, state))
    assert preview["action"] == "local_file_preview"
    assert preview["preview_only"] is True
    assert "before" in preview["preview"]["diff"] and "after" in preview["preview"]["diff"]
    assert target.read_text(encoding="utf-8") == "before\n"

    written = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "write", "path": str(target), "content": "after\n"},
    }, state))
    assert written["undo_available"] is True and written["undo_id"].startswith("undo_")
    assert written["preview"]["changed"] is True
    assert target.read_text(encoding="utf-8") == "after\n"
    undone = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "undo", "undo_id": written["undo_id"]},
    }, state))
    assert undone["restored"] is True
    assert target.read_text(encoding="utf-8") == "before\n"


def test_desktop_patch_rejects_ambiguous_text_and_applies_exact_patch(tmp_path: Path):
    target = tmp_path / "code.py"
    target.write_text("value = 1\nvalue = 1\n", encoding="utf-8")
    state = {"capabilities": ["local_file_write"], "allowed_roots": [str(tmp_path)]}
    with pytest.raises(RuntimeError, match="ambiguous"):
        execute_step({
            "required_capability": "local_file_write",
            "executor_payload": {"operation": "patch", "path": str(target), "find": "value = 1", "replace": "value = 2"},
        }, state)
    result = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "patch", "path": str(target), "find": "value = 1", "replace": "value = 2", "count": 2},
    }, state))
    assert result["operation"] == "patch"
    assert target.read_text(encoding="utf-8") == "value = 2\nvalue = 2\n"


def test_desktop_rename_delete_and_undo_are_bounded_to_approved_root(tmp_path: Path):
    source = tmp_path / "source.txt"
    destination = tmp_path / "renamed.txt"
    source.write_text("keep me", encoding="utf-8")
    state = {"capabilities": ["local_file_write"], "allowed_roots": [str(tmp_path)]}
    moved = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "rename", "path": str(source), "new_path": str(destination)},
    }, state))
    assert not source.exists() and destination.read_text(encoding="utf-8") == "keep me"
    execute_step({"required_capability": "local_file_write", "executor_payload": {"operation": "undo", "undo_id": moved["undo_id"]}}, state)
    assert source.read_text(encoding="utf-8") == "keep me" and not destination.exists()

    deleted = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "delete", "path": str(source)},
    }, state))
    assert not source.exists() and deleted["preview"]["changed"] is True
    execute_step({"required_capability": "local_file_write", "executor_payload": {"operation": "undo", "undo_id": deleted["undo_id"]}}, state)
    assert source.read_text(encoding="utf-8") == "keep me"


def test_desktop_edit_refuses_changed_file_before_undo(tmp_path: Path):
    target = tmp_path / "note.txt"
    target.write_text("one", encoding="utf-8")
    state = {"capabilities": ["local_file_write"], "allowed_roots": [str(tmp_path)]}
    result = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"path": str(target), "content": "two"},
    }, state))
    target.write_text("newer", encoding="utf-8")
    with pytest.raises(RuntimeError, match="changed"):
        execute_step({"required_capability": "local_file_write", "executor_payload": {"operation": "undo", "undo_id": result["undo_id"]}}, state)


def test_desktop_document_toolkit_creates_edits_previews_and_undoes_local_artifacts(tmp_path: Path):
    """Office documents stay inside the existing approved file-write scope."""
    from docx import Document
    from openpyxl import load_workbook
    from pptx import Presentation
    from pypdf import PdfReader

    state = {
        "capabilities": ["local_file_write"],
        "allowed_roots": [str(tmp_path)],
        "_state_path": str(tmp_path / "desktop.json"),
    }
    docx_path = tmp_path / "agent-report.docx"
    preview = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {
            "operation": "create_docx", "path": str(docx_path), "title": "Agent report",
            "sections": [{"heading": "Finding", "body": "Local artifacts are approval-gated.", "bullets": ["No cloud file upload"]}],
            "preview_only": True,
        },
    }, state))
    assert preview["action"] == "local_file_preview" and not docx_path.exists()
    created_docx = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {
            "operation": "create_docx", "path": str(docx_path), "title": "Agent report",
            "sections": [{"heading": "Finding", "body": "Local artifacts are approval-gated."}],
        },
    }, state))
    assert created_docx["document"]["format"] == "docx"
    assert "Agent report" in "\n".join(item.text for item in Document(docx_path).paragraphs)
    edited_docx = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "edit_docx", "path": str(docx_path), "edit": "replace_text", "find": "approval-gated", "replace": "reviewed"},
    }, state))
    assert edited_docx["undo_available"] is True
    assert "reviewed" in "\n".join(item.text for item in Document(docx_path).paragraphs)
    execute_step({"required_capability": "local_file_write", "executor_payload": {"operation": "undo", "undo_id": edited_docx["undo_id"]}}, state)
    assert "approval-gated" in "\n".join(item.text for item in Document(docx_path).paragraphs)

    xlsx_path = tmp_path / "metrics.xlsx"
    created_xlsx = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "create_xlsx", "path": str(xlsx_path), "sheets": [{"name": "Metrics", "freeze_header": True, "rows": [["Metric", "Value"], ["Tasks", 3]]}]},
    }, state))
    assert created_xlsx["document"]["cells"] == 4
    edited_xlsx = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "edit_xlsx", "path": str(xlsx_path), "cells": [{"sheet": "Metrics", "cell": "B3", "formula": "=SUM(B2:B2)"}]},
    }, state))
    assert edited_xlsx["document"]["cells_updated"] == 1
    assert load_workbook(xlsx_path, data_only=False)["Metrics"]["B3"].value == "=SUM(B2:B2)"

    pptx_path = tmp_path / "briefing.pptx"
    created_pptx = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "create_pptx", "path": str(pptx_path), "title": "Smara briefing", "slides": [{"title": "Status", "bullets": ["Local only", "Approval-gated"]}]},
    }, state))
    assert created_pptx["document"]["slides"] == 2
    assert len(Presentation(pptx_path).slides) == 2

    pdf_path = tmp_path / "report.pdf"
    created_pdf = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "create_pdf", "path": str(pdf_path), "title": "Local report", "sections": [{"heading": "Summary", "body": "This PDF stays in the approved workspace."}]},
    }, state))
    assert created_pdf["document"]["format"] == "pdf"
    assert "Local report" in (PdfReader(pdf_path).pages[0].extract_text() or "")
    relative_pdf = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "create_pdf", "path": "relative-report.pdf", "title": "Relative local report"},
    }, state))
    assert relative_pdf["file_name"] == "relative-report.pdf" and (tmp_path / "relative-report.pdf").exists()
    second_pdf = tmp_path / "second.pdf"
    execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "create_pdf", "path": str(second_pdf), "title": "Second local report"},
    }, state)
    merged_pdf = tmp_path / "merged.pdf"
    merged = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "merge_pdf", "path": str(merged_pdf), "source_paths": [str(pdf_path), str(second_pdf)]},
    }, state))
    assert merged["document"]["pages"] == 2 and len(PdfReader(merged_pdf).pages) == 2
    extracted_pdf = tmp_path / "first-page.pdf"
    extracted = json.loads(execute_step({
        "required_capability": "local_file_write",
        "executor_payload": {"operation": "extract_pdf_pages", "path": str(extracted_pdf), "source_path": str(merged_pdf), "pages": [1]},
    }, state))
    assert extracted["document"]["pages"] == [1] and len(PdfReader(extracted_pdf).pages) == 1


def test_desktop_document_toolkit_rejects_wrong_formats_and_unsafe_formulas(tmp_path: Path):
    state = {"capabilities": ["local_file_write"], "allowed_roots": [str(tmp_path)]}
    with pytest.raises(RuntimeError, match="destination path"):
        execute_step({"required_capability": "local_file_write", "executor_payload": {"operation": "create_pdf", "path": str(tmp_path / "wrong.txt"), "title": "No"}}, state)
    with pytest.raises(RuntimeError, match="formulas require"):
        execute_step({
            "required_capability": "local_file_write",
            "executor_payload": {"operation": "create_xlsx", "path": str(tmp_path / "unsafe.xlsx"), "sheets": [{"name": "Sheet1", "rows": [["=WEBSERVICE(\"https://example.com\")"]]}]},
        }, state)


def test_desktop_workspace_inspection_lists_only_bounded_approved_files(tmp_path: Path):
    (tmp_path / "src").mkdir()
    (tmp_path / "src" / "app.py").write_text("print('smara')\n", encoding="utf-8")
    (tmp_path / "README.md").write_text("workspace smoke\n", encoding="utf-8")
    state = {"capabilities": ["local_file_read"], "allowed_roots": [str(tmp_path)]}

    result = json.loads(execute_step({
        "required_capability": "local_file_read",
        "executor_payload": {"operation": "list_tree", "path": str(tmp_path), "max_depth": 3, "max_entries": 10},
    }, state))

    assert result["action"] == "local_workspace_inspect"
    assert result["operation"] == "list_tree"
    assert {item["path"] for item in result["entries"]} >= {"src", "src/app.py", "README.md"}
    assert result["truncated"] is False


def test_desktop_workspace_search_is_bounded_and_text_only(tmp_path: Path):
    (tmp_path / "notes.txt").write_text("First Smara note\nsecond line\n", encoding="utf-8")
    (tmp_path / "binary.bin").write_bytes(b"\x00Smara should not be read")
    state = {"capabilities": ["local_file_read"], "allowed_roots": [str(tmp_path)]}

    result = json.loads(execute_step({
        "required_capability": "local_file_read",
        "executor_payload": {"operation": "search_text", "path": str(tmp_path), "query": "smara", "max_files": 10, "max_matches": 10},
    }, state))

    assert result["action"] == "local_workspace_inspect"
    assert result["scanned_files"] == 2
    assert result["matches"] == [{"path": "notes.txt", "line": 1, "preview": "First Smara note"}]
    with pytest.raises(RuntimeError, match="query"):
        execute_step({
            "required_capability": "local_file_read",
            "executor_payload": {"operation": "search_text", "path": str(tmp_path), "query": ""},
        }, state)


def test_desktop_workspace_filename_search_is_bounded(tmp_path: Path):
    (tmp_path / "notes.txt").write_text("one", encoding="utf-8")
    (tmp_path / "notes.md").write_text("two", encoding="utf-8")
    (tmp_path / "other.py").write_text("three", encoding="utf-8")
    state = {"capabilities": ["local_file_read"], "allowed_roots": [str(tmp_path)]}
    result = json.loads(execute_step({
        "required_capability": "local_file_read",
        "executor_payload": {"operation": "find_files", "path": str(tmp_path), "query": "NOTES", "max_matches": 1},
    }, state))
    assert result["operation"] == "find_files"
    assert len(result["matches"]) == 1 and result["matches"][0] in {"notes.txt", "notes.md"}
    assert result["truncated"] is True


def test_desktop_terminal_requires_allowlist_and_rejects_shell_operators(tmp_path: Path):
    state = {"capabilities": ["local_terminal"], "allowed_roots": [str(tmp_path)], "terminal_allowlist": ["python"]}
    result = json.loads(execute_step({"required_capability": "local_terminal", "executor_payload": {"argv": ["python", "-c", "print(2+2)"], "cwd": str(tmp_path)}}, state))
    assert result["exit_code"] == 0
    assert "4" in result["output"]
    with pytest.raises(RuntimeError, match="Shell operators"):
        execute_step({"required_capability": "local_terminal", "executor_payload": {"command": "python -c \"print(1)\" & whoami", "cwd": str(tmp_path)}}, state)


def test_desktop_named_recipe_reports_artifact_metadata_and_rejects_unknown_recipe(tmp_path: Path):
    (tmp_path / "sample.py").write_text("value = 2\n", encoding="utf-8")
    report = tmp_path / "report.txt"
    report.write_text("local report\n", encoding="utf-8")
    state = {"capabilities": ["local_terminal"], "allowed_roots": [str(tmp_path)], "terminal_allowlist": ["python"]}
    result = json.loads(execute_step({
        "required_capability": "local_terminal",
        "executor_payload": {"recipe": "python.compile", "cwd": str(tmp_path), "artifact_paths": [str(report)]},
    }, state))
    assert result["recipe"] == "python.compile"
    assert result["exit_code"] == 0
    assert result["artifacts"] == [{"path": "report.txt", "bytes": report.stat().st_size, "sha256": result["artifacts"][0]["sha256"]}]
    assert result["changed_files_available"] is False
    with pytest.raises(RuntimeError, match="Unknown local recipe"):
        execute_step({
            "required_capability": "local_terminal",
            "executor_payload": {"recipe": "python.deploy", "cwd": str(tmp_path)},
        }, state)


def test_desktop_recipe_collects_changed_files_when_git_is_allowlisted(tmp_path: Path):
    subprocess.run(["git", "init", "--quiet"], cwd=tmp_path, check=True, capture_output=True)
    (tmp_path / "sample.py").write_text("print('recipe')\n", encoding="utf-8")
    state = {"capabilities": ["local_terminal"], "allowed_roots": [str(tmp_path)], "terminal_allowlist": ["python", "git"]}
    result = json.loads(execute_step({
        "required_capability": "local_terminal",
        "executor_payload": {"recipe": "python.compile", "cwd": str(tmp_path)},
    }, state))
    assert result["changed_files_available"] is True
    assert any(path.startswith("__pycache__/") for path in result["changed_files"])


def test_desktop_terminal_observes_cancellation_before_completion(tmp_path: Path):
    state = {"capabilities": ["local_terminal"], "allowed_roots": [str(tmp_path)], "terminal_allowlist": ["python"]}
    progress: list[str] = []
    with pytest.raises(ExecutionCancelled, match="cancelled"):
        execute_step({
            "required_capability": "local_terminal",
            "executor_payload": {"argv": ["python", "-c", "import time; time.sleep(5)"], "cwd": str(tmp_path)},
        }, state, checkpoint=lambda: True, progress_hook=progress.append)
    assert progress == ["Terminal started: python"]


def test_local_credential_vault_injects_only_requested_alias_and_redacts_output(monkeypatch, tmp_path: Path):
    vault = tmp_path / "credentials.json"
    monkeypatch.setenv("SMARA_DESKTOP_CREDENTIALS", str(vault))
    secret = "local-only-test-secret"
    save_local_credential("TAVILY_API_KEY", secret, "tavily")
    assert secret not in vault.read_text(encoding="utf-8") if os.name == "nt" else True
    assert local_credential_summaries()[0]["name"] == "TAVILY_API_KEY"
    assert resolve_local_credential("TAVILY_API_KEY") == secret
    state = {"capabilities": ["local_terminal"], "allowed_roots": [str(tmp_path)], "terminal_allowlist": ["python"]}
    result = json.loads(execute_step({
        "required_capability": "local_terminal",
        "executor_payload": {
            "argv": ["python", "-c", "import os; print(os.environ['TAVILY_API_KEY'])"],
            "cwd": str(tmp_path),
            "credential_env": ["TAVILY_API_KEY"],
        },
    }, state))
    assert secret not in result["output"]
    assert "REDACTED LOCAL CREDENTIAL" in result["output"]
    assert result["credential_env"] == ["TAVILY_API_KEY"]
    assert delete_local_credential("TAVILY_API_KEY") is True
    assert local_credential_summaries() == []


def test_local_tavily_adapter_uses_vault_and_returns_bounded_secret_free_proof(monkeypatch, tmp_path: Path):
    vault = tmp_path / "credentials.json"
    monkeypatch.setenv("SMARA_DESKTOP_CREDENTIALS", str(vault))
    secret = "tavily-local-secret"
    save_local_credential("TAVILY_API_KEY", secret, "tavily")
    calls: list[dict] = []

    class Response:
        status_code = 200
        is_success = True

        def json(self):
            return {"results": [{"title": "Smara", "url": "https://example.com/smara", "content": "A bounded result."}]}

    class Client:
        def __init__(self, **_kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def post(self, url, **kwargs):
            calls.append({"url": url, **kwargs})
            return Response()

    monkeypatch.setattr("smara.desktop_executor.httpx.Client", Client)
    state = {"capabilities": ["local_integration"], "allowed_roots": [str(tmp_path)]}
    result = json.loads(execute_step({
        "required_capability": "local_integration",
        "executor_payload": {"provider": "tavily", "operation": "search", "query": "Smara"},
    }, state))
    assert calls[0]["url"] == "https://api.tavily.com/search"
    assert calls[0]["json"]["api_key"] == secret
    assert secret not in json.dumps(result)
    assert result["results"][0]["url"] == "https://example.com/smara"
    assert result["citations"] == ["https://example.com/smara"]
    assert result["proof"]["results"] == 1
    assert result["connector"] == {
        "provider": "tavily", "operation": "search", "auth_mode": "local_api_key", "risk": "read_only",
        "scopes": ["web.search"], "max_results": 5, "max_requests_per_run": 1,
    }


def test_local_github_adapter_is_read_only_and_classifies_bad_credentials(monkeypatch, tmp_path: Path):
    vault = tmp_path / "credentials.json"
    monkeypatch.setenv("SMARA_DESKTOP_CREDENTIALS", str(vault))
    save_local_credential("GITHUB_TOKEN", "github-local-secret", "github")

    class Response:
        status_code = 200
        is_success = True

        def json(self):
            return [{"name": "smara", "full_name": "sujal/smara", "private": True, "html_url": "https://github.com/sujal/smara", "description": "agent", "language": "Python"}]

    class Client:
        def __init__(self, **_kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def get(self, url, **_kwargs):
            assert url == "https://api.github.com/user/repos"
            return Response()

    monkeypatch.setattr("smara.desktop_executor.httpx.Client", Client)
    state = {"capabilities": ["local_integration"], "allowed_roots": [str(tmp_path)]}
    result = json.loads(execute_step({
        "required_capability": "local_integration",
        "executor_payload": {"provider": "github", "operation": "list_repositories", "limit": 1},
    }, state))
    assert result["repositories"][0]["full_name"] == "sujal/smara"
    assert "github-local-secret" not in json.dumps(result)
    assert result["connector"]["scopes"] == ["repositories:read"]


def test_local_connector_catalogue_is_explicit_and_secret_free():
    connectors = {item["provider"]: item for item in local_connector_catalog()}
    assert set(connectors) == {"tavily", "github"}
    assert connectors["tavily"]["credential_alias"] == "TAVILY_API_KEY"
    assert connectors["github"]["operation"] == "list_repositories"
    assert all(item["credential_configured"] is False for item in connectors.values())


def test_local_connector_lifecycle_records_only_proof_and_can_revoke(monkeypatch, tmp_path: Path):
    vault = tmp_path / "credentials.json"
    audit = tmp_path / "connector-audit.json"
    monkeypatch.setenv("SMARA_DESKTOP_CREDENTIALS", str(vault))
    monkeypatch.setenv("SMARA_DESKTOP_CONNECTOR_AUDIT", str(audit))
    secret = "never-write-this-to-the-audit"
    save_local_credential("TAVILY_API_KEY", secret, "tavily")

    class Response:
        status_code = 200
        is_success = True
        def json(self): return {"results": [{"title": "Source", "url": "https://example.com", "content": "private query output"}]}

    class Client:
        def __init__(self, **_kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def post(self, *_args, **_kwargs): return Response()

    monkeypatch.setattr("smara.desktop_executor.httpx.Client", Client)
    summaries = {item["provider"]: item for item in local_connector_summaries()}
    assert summaries["tavily"]["credential_configured"] is True
    result = execute_step({
        "required_capability": "local_integration",
        "executor_payload": {"provider": "tavily", "operation": "search", "query": "private search phrase"},
    }, {"capabilities": ["local_integration"], "allowed_roots": [str(tmp_path)]})
    events = local_connector_audit()
    assert events[-1]["provider"] == "tavily"
    assert events[-1]["status"] == "completed"
    assert events[-1]["proof"]["results"] == 1
    serialized = json.dumps(events)
    assert secret not in serialized and "private search phrase" not in serialized and "private query output" not in serialized
    assert "result_sha256" in serialized and result
    assert revoke_local_connector("tavily") is True
    assert local_connector_audit()[-1]["status"] == "revoked"
    assert {item["provider"]: item for item in local_connector_summaries()}["tavily"]["credential_configured"] is False


def test_unapproved_local_connector_step_never_calls_provider(monkeypatch, tmp_path: Path):
    called = False

    def unexpected(*_args, **_kwargs):
        nonlocal called
        called = True
        raise AssertionError("provider must not run before approval")

    monkeypatch.setattr("smara.desktop_executor.execute_local_integration", unexpected)
    with pytest.raises(RuntimeError, match="approval"):
        execute_step({
            "requires_approval": True,
            "required_capability": "local_integration",
            "executor_payload": {"provider": "tavily", "operation": "search", "query": "safe"},
        }, {"capabilities": ["local_integration"], "allowed_roots": [str(tmp_path)]})
    assert called is False


def test_desktop_refuses_unapproved_step_and_undeclared_browser(tmp_path: Path):
    state = {"capabilities": ["local_browser"], "allowed_roots": [str(tmp_path)], "browser_domains": ["example.com"]}
    with pytest.raises(RuntimeError, match="approval"):
        execute_step({"requires_approval": True, "required_capability": "local_browser", "executor_payload": {"url": "https://example.com"}}, state)
    with pytest.raises(RuntimeError, match="not declared"):
        execute_step({"required_capability": "local_file_read", "executor_payload": {"path": str(tmp_path / "x")}}, state)


def test_desktop_browser_inspection_is_text_only_and_has_no_browser_session(monkeypatch, tmp_path: Path):
    class Response:
        is_redirect = False
        headers = {"content-type": "text/html; charset=utf-8"}
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def raise_for_status(self): return None
        def iter_bytes(self): yield b"<html><title>Local test</title><body>Hello <b>Smara</b><script>secret()</script></body></html>"

    class Client:
        def __init__(self, **_kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def stream(self, method, url):
            assert method == "GET" and url == "https://example.com/docs"
            return Response()

    monkeypatch.setattr("smara.desktop_executor.httpx.Client", Client)
    state = {"capabilities": ["local_browser"], "allowed_roots": [str(tmp_path)], "browser_domains": ["example.com"]}
    result = json.loads(execute_step({
        "required_capability": "local_browser",
        "executor_payload": {"operation": "inspect_text", "url": "https://example.com/docs"},
    }, state))
    assert result["operation"] == "inspect_text"
    assert result["title"] == "Local test"
    assert "Hello" in result["text"] and "Smara" in result["text"]
    assert "secret" not in result["text"]


def test_desktop_browser_dom_inspection_is_bounded_and_resolves_links(monkeypatch, tmp_path: Path):
    class Response:
        is_redirect = False
        headers = {"content-type": "text/html; charset=utf-8", "content-length": "280"}
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def raise_for_status(self): return None
        def iter_bytes(self):
            yield (
                b"<html><title>Controls</title><body><main id='app'>"
                b"<h1>Welcome</h1><button class='primary' aria-label='go'>"
                b"Go <span>now</span></button><a href='/docs'>Docs</a>"
                b"<script><button>secret</button></script></main></body></html>"
            )

    class Client:
        def __init__(self, **_kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def stream(self, method, url):
            assert method == "GET" and url == "https://example.com/app"
            return Response()

    monkeypatch.setattr("smara.desktop_executor.httpx.Client", Client)
    state = {"capabilities": ["local_browser"], "allowed_roots": [str(tmp_path)], "browser_domains": ["example.com"]}
    result = json.loads(execute_step({
        "required_capability": "local_browser",
        "executor_payload": {"operation": "inspect_dom", "url": "https://example.com/app", "selector": "button.primary", "max_elements": 1},
    }, state))
    assert result["title"] == "Controls"
    assert result["count"] == 1
    assert result["elements"][0]["tag"] == "button"
    assert result["elements"][0]["text"] == "Go now"
    assert result["elements"][0]["attributes"]["aria-label"] == "go"
    assert "secret" not in json.dumps(result)
    assert len(result["proof"]["content_sha256"]) == 64


def test_desktop_browser_download_is_atomic_bounded_and_scoped(monkeypatch, tmp_path: Path):
    download_dir = tmp_path / "downloads"
    download_dir.mkdir()

    class Response:
        is_redirect = False
        headers = {"content-type": "application/octet-stream", "content-length": "11"}
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def raise_for_status(self): return None
        def iter_bytes(self):
            yield b"hello "
            yield b"smara"

    class Client:
        def __init__(self, **_kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def stream(self, method, url):
            assert method == "GET" and url == "https://example.com/file.bin"
            return Response()

    monkeypatch.setattr("smara.desktop_executor.httpx.Client", Client)
    state = {"capabilities": ["local_browser"], "allowed_roots": [str(tmp_path)], "browser_domains": ["example.com"]}
    step = {
        "required_capability": "local_browser",
        "executor_payload": {
            "operation": "download", "url": "https://example.com/file.bin", "destination": "downloads/file.bin",
        },
    }
    result = json.loads(execute_step(step, state))
    target = download_dir / "file.bin"
    assert target.read_bytes() == b"hello smara"
    assert result["path"] == "downloads/file.bin"
    assert result["bytes_downloaded"] == 11
    assert result["overwrote"] is False
    with pytest.raises(RuntimeError, match="already exists"):
        execute_step(step, state)
    step["executor_payload"]["overwrite"] = True
    replaced = json.loads(execute_step(step, state))
    assert replaced["overwrote"] is True


def test_desktop_browser_download_cancellation_removes_partial_file(monkeypatch, tmp_path: Path):
    class Response:
        is_redirect = False
        headers = {"content-type": "application/octet-stream"}
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def raise_for_status(self): return None
        def iter_bytes(self): yield b"partial"

    class Client:
        def __init__(self, **_kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *_args): return None
        def stream(self, _method, _url): return Response()

    monkeypatch.setattr("smara.desktop_executor.httpx.Client", Client)
    state = {"capabilities": ["local_browser"], "allowed_roots": [str(tmp_path)], "browser_domains": ["example.com"]}
    with pytest.raises(ExecutionCancelled, match="cancelled"):
        execute_step({
            "required_capability": "local_browser",
            "executor_payload": {"operation": "download", "url": "https://example.com/file.bin", "destination": "file.bin"},
        }, state, checkpoint=lambda: True)
    assert not (tmp_path / "file.bin").exists()


def test_desktop_state_round_trip_is_json(tmp_path: Path):
    path = tmp_path / "desktop.json"
    _save_state(path, {"executor_id": "desktop_1", "token": "opaque", "smara_url": "http://smara"})
    assert _load_state(path)["executor_id"] == "desktop_1"
    if os.name == "nt":
        stored = path.read_text(encoding="utf-8")
        assert "opaque" not in stored
        assert "token_dpapi" in stored


def test_desktop_runner_refreshes_step_before_completion(monkeypatch, tmp_path: Path):
    calls = []
    step = {
        "step_id": "step_1",
        "required_capability": "local_file_read",
        "executor_payload": {"path": str(tmp_path / "note.txt")},
    }

    class Response:
        status_code = 200

        def __init__(self, payload=None):
            self.payload = payload

        def json(self):
            return self.payload or {}

        def raise_for_status(self):
            return None

    class Client:
        def post(self, path, **kwargs):
            calls.append(path)
            if path.endswith("/claim"):
                return Response({"step": step})
            return Response()

    monkeypatch.setattr("smara.desktop_executor.execute_step", lambda *_args, **_kwargs: "result")
    state = {
        "smara_url": "https://smara.example",
        "executor_id": "desktop_1",
        "token": "opaque",
        "capabilities": ["local_file_read"],
    }
    state_path = tmp_path / "desktop.json"
    state_path.write_text(json.dumps(state), encoding="utf-8")
    assert DesktopRunner(state_path).run_once(Client(), state) is True
    assert calls == [
        "https://smara.example/v1/executors/heartbeat",
        "https://smara.example/v1/executors/claim",
        "https://smara.example/v1/executors/steps/step_1/heartbeat",
        "https://smara.example/v1/executors/steps/step_1/complete",
    ]
